
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION, WD_ORIENTATION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_font(run, font_name='Times New Roman', font_size=12, bold=False, italic=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Necessary for TNR to apply to all characters in some versions
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_header_footer(doc, project_id, college_name):
    for section in doc.sections:
        # Header
        header = section.header
        htable = header.add_table(1, 2, Inches(6.25))
        htable.columns[0].width = Inches(3.125)
        htable.columns[1].width = Inches(3.125)
        
        # Top Left: Project ID
        p_left = htable.cell(0, 0).paragraphs[0]
        p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_left = p_left.add_run(f"Project ID: {project_id}")
        set_font(run_left, font_size=10)
        
        # Top Right: Chapter Heading (Placeholder logic)
        p_right = htable.cell(0, 1).paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_right = p_right.add_run("Chapter Heading Replacement") # Dynamic would be better but static for now
        set_font(run_right, font_size=10)

        # Footer
        footer = section.footer
        ftable = footer.add_table(1, 3, Inches(6.25))
        ftable.columns[0].width = Inches(2.08)
        ftable.columns[1].width = Inches(2.08)
        ftable.columns[2].width = Inches(2.08)
        
        # Bot Left: GTU
        p_left = ftable.cell(0, 0).paragraphs[0]
        p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_left = p_left.add_run("Gujarat Technological University")
        set_font(run_left, font_size=10)
        
        # Bot Right: College
        p_right = ftable.cell(0, 2).paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_right = p_right.add_run(college_name)
        set_font(run_right, font_size=10)
        
        # Bot Center: Page Number
        # (Page numbering in python-docx is tricky, usually requires complex XML or field codes)
        # We will handle it by just placing a placeholder or using a known method if possible
        p_center = ftable.cell(0, 1).paragraphs[0]
        p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # run_center = p_center.add_run("Page | ") # Page numbering usually requires Word to calculate

def add_section(doc, title, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(title)
    if level == 1:
        set_font(run, font_size=16, bold=True) # Chapter Heading
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif level == 2:
        set_font(run, font_size=14, bold=True) # Section Heading
    else:
        set_font(run, font_size=12, bold=True) # Subsection

def create_report():
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.page_height = Inches(11.69) # A4
    section.page_width = Inches(8.27)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    
    # Data
    student_name = "Dhruv Patni"
    enrollment = "230023107054"
    guide_name = "Ms. Rupal Shah"
    college_name = "Ahmedabad Institute of Technology"
    company_name = "GTCSYS Services Private Limited"
    project_title = "LaraBids - A Real-Time Premium Auction Ecosystem"
    project_id = "[LaraBids-2026-ID]"
    
    # Cover Page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n\n\nInternship Project Report on")
    set_font(run, font_size=16, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n{project_title}")
    set_font(run, font_size=20, bold=True)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n\n\nSubmitted by:\n{student_name}\n({enrollment})")
    set_font(run, font_size=14, bold=True)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n\n\nInternal Guide:\n{guide_name}")
    set_font(run, font_size=14, bold=True)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n\n\nUnder Internship at:\n{company_name}")
    set_font(run, font_size=14, bold=True)
    
    doc.add_page_break()
    
    # 2. Certificates (College)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("COLLEGE CERTIFICATE")
    set_font(run, font_size=14, bold=True)
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(f"\nThis is to certify that Mr. {student_name}, Enrollment No: {enrollment}, of Computer Engineering department, 8th Semester has successfully completed the Industry Internship project entitled \"{project_title}\" at {company_name} during the academic year 2025-26.")
    set_font(run, font_size=14)
    
    doc.add_page_break()
    
    # 3. Company Certificate
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("COMPANY CERTIFICATE")
    set_font(run, font_size=14, bold=True)
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(f"\nTo Whom It May Concern,\nThis is to certify that Mr. {student_name} has completed his internship at {company_name}. During his tenure, he worked on the \"LaraBids\" project. His performance was exceptional, demonstrating strong technical skills in Laravel, MySQL, and real-time system architecture.")
    set_font(run, font_size=14)
    
    doc.add_page_break()
    
    # 4. Declaration
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CANDIDATE’S DECLARATION")
    set_font(run, font_size=14, bold=True)
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(f"\nI, {student_name}, hereby declare that the project report entitled \"{project_title}\" is a result of my own architectural design and implementation work carried out at {company_name} under the guidance of Mr. Mukesh Lagadhir (Industry Mentor) and {guide_name} (Internal Guide).")
    set_font(run, font_size=14)
    
    doc.add_page_break()
    
    # 5. Acknowledgement
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ACKNOWLEDGEMENT")
    set_font(run, font_size=14, bold=True)
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run("\nI wish to express my deep sense of gratitude to GTCSYS Services Private Limited and specifically to Mr. Mukesh Lagadhir for providing me with the opportunity to work on such a high-impact project. I am also thankful to Ms. Rupal Shah, my internal guide at Ahmedabad Institute of Technology, for her constant support and valuable insights throughout the internship period.")
    set_font(run, font_size=12)
    
    doc.add_page_break()
    
    # 6. Abstract
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ABSTRACT")
    set_font(run, font_size=14, bold=True)
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run("\nLaraBids is a modern, high-performance web-based auction platform designed to bridge the gap between traditional manual bidding and high-scale digital commerce. Built using the Laravel 12 framework, the system provides a 'premium-feel' user experience while ensuring robust transactional integrity and real-time responsiveness. The core objective of LaraBids is to eliminate physical geographical barriers and lack of transparency in the auction industry. The system implements advanced bidding mechanics, including Proxy Bidding and a Real-Time Bidding Engine built on MySQL transactions.")
    set_font(run, font_size=14, italic=True)
    
    doc.add_page_break()
    
    # --- Chapters ---
    def read_draft():
        try:
            with open(r'c:\LaraBids\Project_Report_Draft.md', 'r', encoding='utf-8') as f:
                return f.read()
        except:
            return ""

    draft_content = read_draft()
    
    current_chapter = ""
    current_section = ""
    
    lines = draft_content.split('\n')
    for line in lines:
        if line.startswith('## '):
            current_chapter = line.replace('## ', '').strip()
            add_section(doc, current_chapter, 1)
        elif line.startswith('### '):
            current_section = line.replace('### ', '').strip()
            add_section(doc, current_section, 2)
        elif line.strip() and not line.startswith('#'):
            # Text content
            p = doc.add_paragraph(line.strip())
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if len(p.runs) > 0:
                set_font(p.runs[0], font_size=12)
            else:
                run = p.add_run()
                set_font(run, font_size=12)
    
    # Add Header/Footer
    add_header_footer(doc, project_id, college_name)
    
    doc.save("Internship_Project_Report_Final.docx")
    print("Report generated: Internship_Project_Report_Final.docx")

if __name__ == "__main__":
    create_report()
